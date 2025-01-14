from collections import defaultdict
from datetime import datetime
from typing import Dict


import pymupdf, os
from docx import Document
from docx.shared import Inches

class DocumentHandler:
    def __init__(self, path:str, content: str, title: str):
        self.path = path
        self.document_type = self.get_document_type(self.path)
        self.content = content
        self.content_tokens = self.content.split(" ")
        self.author = "unknown"
        self.title = title
        self.file_name = f"{self.author}-{self.title}" if self.author != "unknown" else self.title
        self.short_summary = ""
        self.summary = ""
        self.sentiment = 0.0
        self.sentiment_reason = str
        self.entities = []
        self.marked_sentences = defaultdict(list) # Initialize dictionary to store sentences by marking color
        self.wordcloud_data = dict() # Dict to store the information for the wordclouds
        self.answers = dict() # Dict where the answer dict (quesiton/reasoning/answer) of every question is stored
        self.topic_clusters = dict() # Dict where the cluster name and the list of topics is stored

        print(f'{datetime.now().strftime("%H:%M:%S")} Initialized Document: <{self.title}>')


    def get_document_type(self, path: str):
        return path.split(".")[-1]
    
    def __str__(self) -> str:
        """ Returns a short summary if called. """

        return_string = f"Title: {self.title}\n"
        return_string += f"Short Summary: {self.short_summary}\n"
            
        return return_string
            
            
    def get_info(self):
        """ Extends __str__ method with further information. """
        return_string = str(self)
        return_string += f"Summary:\n{self.summary}\n"
        return_string += f"Sentiment: {self.sentiment}\n"
        return_string += f"Entities: {self.entities}\n"
        return_string += f"Topic clusters: {self.topic_clusters}\n"
        return_string += f"Word Cloud Data: {self.wordcloud_data}\n"
            
        for question, answer in self.answers.items():
            return_string += f"Question: {question}\n\t{answer}\n"
        
        return return_string
           
        
    def to_dict(self) -> Dict:
        """ Returns the instance as dict (e.g., for saving). """

        return {
            "path": self.path,
            "content": self.content,
            "content_tokens": self.content_tokens,
            "author": self.author,
            "title": self.title,
            "short_summary": self.short_summary,
            "summary": self.summary,
            "sentiment": self.sentiment,
            "sentiment_reason": self.sentiment_reason,
            "entities": self.entities,
            "marked_sentences": self.marked_sentences,
            "answers": self.answers,
            "topic_clusters": self.topic_clusters,
            "wordcloud_data": self.wordcloud_data
        }
    
    @classmethod
    def from_dict(cls, data: Dict):
        """ Loads an instance of this class based on to_dict() method. """

        doc = cls(data.get("path"), data.get("content"), data.get("title"))
        doc.document_type = doc.get_document_type(doc.path)
        doc.content_tokens = data.get("content_tokens")
        doc.author = data.get("author")
        doc.short_summary = data.get("short_summary")
        doc.summary = data.get("summary")
        doc.sentiment = data.get("sentiment")
        doc.sentiment_reason = data.get("sentiment_reason")
        doc.entities = data.get("entities")
        doc.marked_sentences = data.get("marked_sentences")
        doc.answers = data.get("answers")
        doc.topic_clusters = data.get("topic_clusters")
        doc.wordcloud_data = data.get("wordcloud_data")
        return doc
            
    def get_markdown(self) -> str:
        """Formats document data into a markdown string. """

        # Handle potential errors gracefully
        try:
            title = self.title
        except AttributeError:
            title = "No Title"
            
        try:
            short_summary = self.short_summary
        except AttributeError:
            short_summary = "No Short Summary"
            
        try:
            summary = self.summary
        except AttributeError:
            summary = "No Summary"
        
        try:
            questions = ""
            for id, (question, answer) in enumerate(self.answers.items()):
                questions += f"## Question {id+1}:\n"
                questions += f"*{question}*\n\n"
                questions += f"{answer}\n"
                
        except AttributeError:
            questions = "No Questions"
            
        try:
            sentiment = self.sentiment
        except AttributeError:
            sentiment = "No Sentiment"
            
        try:
            entities = self.entities
            entities_str = str(entities)
        except AttributeError:
            entities_str = "No Entities"
            
        try:
            topics = self.topic_clusters
            topics_str = "".join([f"## {cluster}\n\t{', '.join(topics)}\n" for cluster, topics in topics.items()])
        except AttributeError:
            topics_str = "No Topic Clusters"

        # Construct the markdown string
        markdown_output = f"# Media Analysis - {title}\n\n"
        markdown_output += f"# Short Summary\n{short_summary}\n\n"
        markdown_output += f"# Summary\n{summary}\n\n"
        markdown_output += f"# Questions\n{questions}\n"
        markdown_output += f"# Sentiment\n{sentiment}\n\n"
        markdown_output += f"# Entities\n{entities_str}\n\n"
        markdown_output += f"# Topic Clusters\n{topics_str}\n"
        #markdown_output += f"# Markings\n{markings}\n\n"

        # No word cloud part because it's difficult to show it in markdown
        
        return markdown_output
    
    def save_as_docx(self, file_path: str):
        """Exports document data, including word cloud, to a Word document."""
        docx = Document()

        # Add content to the Word document
        docx.add_heading(f"Media Analysis - {self.title}", level=1)

        # Summaries
        docx.add_heading("Short Summary", level=2)
        docx.add_paragraph(self.short_summary)

        docx.add_heading("Summary", level=2)
        docx.add_paragraph(self.summary)
                    
        # Questions
        for id, (question, answer) in enumerate(self.answers.items()):
            docx.add_heading(f"Question {id+1}:", level=2)
            docx.add_paragraph(question)
            docx.add_paragraph(answer)

        # Sentiment
        docx.add_heading("Sentiment", level=2)
        docx.add_paragraph(f"The sentiment is <{self.sentiment}>")
        
        ## Get sentiment reason or first reason if it is a list
        if isinstance(self.sentiment_reason, list):
            docx.add_paragraph(f"First reason for sentiment: {self.sentiment_reason[0]}")
        else:
            docx.add_paragraph(f"{self.sentiment_reason}")

        # Entities
        docx.add_heading("Entities", level=2)
        docx.add_paragraph(", ".join(self.entities))
        
        # Topics
        doc_topics = self.topic_clusters.items()

        if doc_topics is not None:
            docx.add_heading("Topic Clusters", level=2)
            for cluster, topics in doc_topics:
                docx.add_heading(f"Cluster: {cluster}", level=3)
                docx.add_paragraph(", ".join(topics))

        # Word Cloud        
        docx.add_heading("Word Clouds", level=2)
        for key, values in self.wordcloud_data.items():
            try:
                (_, word_frequencies), (_, wordcloud_path) = values.items()
                
                if key == "markings" and values:    
                    num_markings = self.get_number_of_markings()
                    if num_markings is not None:
                        docx.add_heading("Markings", level=2)
                        for color, count in num_markings.items():
                            docx.add_paragraph(f"{color}: {count}")
                    
                
                # Check if the word cloud image exists
                if os.path.exists(wordcloud_path):  
                    docx.add_heading(f"Wordcloud of {key}", level=3)
                    docx.add_picture(wordcloud_path, width=Inches(4.0))
                
                    if word_frequencies:    
                        docx.add_paragraph("Top 10 words:") 
                        # Add top words and their frequencies
                        word_freqs = [f"- {word}: {freq}" for word, freq in word_frequencies]
                        docx.add_paragraph("\n".join(word_freqs))
                        
            except ValueError:
                pass
                        
        # Save the document
        docx.save(file_path)

        
class PdfDocument(DocumentHandler):
    """ Implements function to extract marked text based on text color. """
    def __init__(self, path: str, content: str, title: str):
        super().__init__(path, content, title)
        
        print(f'{datetime.now().strftime("%H:%M:%S")} Initialized PdfDocument: <{self.title}>')
    
    
    def extract_marked_sentences(self) -> dict:
        
        marked_sentences = defaultdict(list)   

        # Open PDF document
        doc = pymupdf.open(self.path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Get plain text content of the page
            text = page.get_text()
            
            # Split text into sentences (basic splitting by period)
            sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
            
            # Get markings on the page
            markings = page.get_text_words()
            annots = page.annots()
            
            if annots:
                for annot in annots:
                    if annot.type[0] == 8:  # Marking annotation
                        # Get marking coordinates
                        coords = annot.rect
                        
                        # Get color of marking (normalize to RGB)
                        color = annot.colors['stroke']
                        if color:
                            color_rgb = tuple(int(c * 255) for c in color)
                        else:
                            continue
                        
                        # Find words within marking coordinates
                        marked_words = []
                        for word_info in markings:
                            word_rect = pymupdf.Rect(word_info[:4])
                            if coords.intersects(word_rect):
                                marked_words.append(word_info[4])
                        
                        if marked_words:
                            # Find the sentence containing the marked words
                            for sentence in sentences:
                                if any(word.lower() in sentence.lower() for word in marked_words):
                                    # Convert RGB tuple to hex for consistent key format
                                    color_hex = '#{:02x}{:02x}{:02x}'.format(*color_rgb)
                                    if sentence not in marked_sentences[color_hex]:
                                        marked_sentences[color_hex].append(sentence)
        
        # Convert defaultdict to regular dict before returning it to the class
        self.marked_sentences = dict(marked_sentences)

    def get_pretty_markings(self):
        """ Returns a prettified markings for every color and sentence. """
        text = ""
        
        for color, sentences in self.marked_sentences.items():
            print(f"\nMarking Color: {color}")
            print("-" * 50)
            
            text += "".join([f"{i}. {sentence}\n" for i, sentence in enumerate(sentences, 1)])
            
        return text 
        
                
    def get_number_of_markings(self) -> dict:
        """ Function to get the number of markings in the document. """
        colors_count = dict()

        for color, sentences in self.marked_sentences.items():
            colors_count[color] = len(sentences)
            
        return colors_count
        
    

if __name__ == "__main__":
    print("ToDo...")