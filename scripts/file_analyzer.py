from .ollama_handler import OllamaMediaAnalysis
from .wordcloud_handler import WordCloudHandler
from .document_handler import *

import re, glob
from datetime import datetime

from pypdf import PdfReader
from docx import Document

import spacy
import dill as pickle

class FileAnalyzer:
    def __init__(self, ollama_handler: OllamaMediaAnalysis, questions: list[str], file_name: str, output_folder: str = "", entity_collection = "all", spacy_model: str = "en_core_web_sm", get_markings: bool = False, debug=False, speed_debug=False):
        """ Manages all files in the provided folder

        Args:
            ollama_handler (OllamaMediaAnalysis): Ollama Handler for LLM.
            output_folder (str, optional): Output folder of all files. Defaults to "".
            questions (list[str]): Questions that should be processed.
            entity_collection (str, optional): Indicates whether the entities should be extracted with spacy or ollama. Defaults to "all".
            get_markings (bool, optional): Indicates whether the markings (markings) should be exported. Defaults to False.
            debug (bool, optional): Whether outputs should be printed. Defaults to True.
            speed_debug (bool, optional): Whether only summaries (True) or all other questions should be processed. Defaults to False.
        """
        self.ollama_handler = ollama_handler
        self.wordcloud = WordCloudHandler()
        
        try:
            self.nlp = spacy.load(spacy_model)
            
        except OSError:
            spacy.cli.download(spacy_model)
            self.nlp = spacy.load(spacy_model)
    
        self.file_name = file_name
        self.output_folder = output_folder
        self.entitiy_collection = entity_collection if entity_collection in ["all", "ollama", "spacy"] else "all"
        self.all_documents = []
        self.get_markings = get_markings
        self.questions = questions
        self.debug = debug
        self.speed_debug = speed_debug
        self.analysis = dict()
        
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Iterates over all pages in the document and stores the text in instance. """

        reader = PdfReader(pdf_path)
        text = ""
        num_pages = reader.pages
        
        for page_count, page in enumerate(num_pages):
            text_current_page = page.extract_text()
            print(f"{datetime.now().strftime("%H:%M:%S")}\t Adding page {page_count}/{len(num_pages)} with {len(text_current_page)} characters")
            text += text_current_page
        return text
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Iterates over all paragraphs in the .docx document and extracts the text. """
        document = Document(docx_path)
        text = ""
        
        for idx, paragraph in enumerate(document.paragraphs):
            text_current_paragraph = paragraph.text.strip()
            print(f"{datetime.now().strftime('%H:%M:%S')}\t Adding paragraph {idx + 1}/{len(document.paragraphs)} with {len(text_current_paragraph)} characters")
            text += text_current_paragraph + "\n"
        
        return text
    
    def extract_entities(self, text: str) -> list[str]:
        """Extracts entities from text using spacy PERSON and ORG labels. """
        
        doc = self.nlp(text)
        entities = [ent.text for ent in doc.ents if ent.label_ in ["PERSON", "ORG"]]
        print(f"{datetime.now().strftime("%H:%M:%S")}\t Found {len(entities)} in text")
        return list(set(entities))
    
    def get_tokens(self, text) -> str:
        """ Tokenize and remove stop words. """
        cleaned_text = self.clean_input(text, hard_clean=True, line_breaks=True)
        doc = self.nlp(cleaned_text)
        
        tokens = [token.lemma_ for token in doc if not token.is_stop and not token.is_punct and not token.is_currency and not token.is_digit and token.is_alpha]
        tokens = [token.title() for token in tokens if token.isupper() or token.capitalize]
        return ' '.join(tokens)
    
    def process(self, doc: DocumentHandler | PdfDocument):
        """ Main method to process all prompt. Should be called differently based on document object (e.g., pdf or docx) """

        # Generating tokenized content
        print(f"{datetime.now().strftime("%H:%M:%S")}\t Generating tokenized content")
        doc.content_tokens = self.get_tokens(doc.content)
        
        # Generating short summary
        short_summary_response = self.ollama_handler.generate_short_summary(doc.content)
        doc.short_summary = self.clean_input(short_summary_response, hard_clean=False, line_breaks=False)
        
        # Generate long summary        
        summary_response = self.ollama_handler.generate_summary(doc.content)
        doc.summary = self.clean_input(summary_response, hard_clean=False, line_breaks=False)
        
        if self.debug:
            print(f"Types: summary={doc.summary}, short={doc.short_summary}")

        if not self.speed_debug:
            # Get answers to questions
            print(f"{datetime.now().strftime("%H:%M:%S")}\t Finding answer to {len(self.questions)} question{"s" if len(self.questions) > 1 else ""}")
            for question in self.questions:
                question_response = self.ollama_handler.answer_question(doc.content, question)
                question_response = self.clean_input(question_response, hard_clean=False, line_breaks=False)
                
                doc.answers[question] = question_response.get("answer")
                        
            # Get sentiment           
            sentiment_response = self.ollama_handler.analyze_sentiment(doc.content) 
            doc.sentiment = sentiment_response.get("sentiment_value")
            doc.sentiment_reason = sentiment_response.get("sentiment_reason")
            
            # Get entities
            print(f"{datetime.now().strftime("%H:%M:%S")}\t Extracting entities from text")
            if self.entitiy_collection in ["all", "spacy"]:
                # Get entities with spacy
                entities_response = self.extract_entities(doc.content_tokens)
                doc.entities = self.clean_input(entities_response, hard_clean=True, line_breaks=True)
                
            # Get markings
            if self.get_markings:
                print(f"{datetime.now().strftime("%H:%M:%S")}\t Extracting text-markings")
                doc.extract_marked_sentences()
            
            # Get topic clusters
            topics_response = self.ollama_handler.extract_topics(doc.content)
            topic_clusters_response = self.ollama_handler.create_topic_clusters(topics_response)
            
            doc.topic_clusters = topic_clusters_response
        
        # Process wordclouds
        doc = self.create_wordcloud(doc, wordcloud_names=["markings", "content", "summary"])
        
        return doc
    
    def extract_title(self, path):
        # Get the filename (basename) and remove the file property
        title = os.path.splitext(os.path.basename(path))[0]
        
        # Clean from special characters
        title_strings = title.split("_")
        
        if isinstance(title_strings, list):
            title = ""              
            for sub_string in title_strings:
                if len(sub_string) > len(title):
                    title = sub_string
            
        # Removing linebreaks, tabs etc.
        title = self.clean_input(title, hard_clean=False, line_breaks=True)
                
        return title
    
    def process_docx(self, docx_path: str) -> DocumentHandler:
        """ Extracts the content from the docx_path and process all questions

        Returns:
            DocumentHandler: Document Handler with enriched informations.
        """
        
        content = self.extract_text_from_docx(docx_path)
        content = self.clean_input(content, hard_clean=True, line_breaks=True)
        
        # Remove the top boxes with author link etc.
        try:
            new_content = content.split("Main image")
            if len(new_content[1]) > len(new_content[0]):
                # Set new content to part after main image
                content = new_content[1]
        except Exception as e:
            print(f"ERROR: Could not remove article properties")
            new_content = content
        
        # Extract author and remove it
        try: 
            # Set author
            author = new_content[0].split("Author ")[1]
            author = self.clean_input(author)
            
            # Replace the surname of the author in the text with a placeholder
            content = content.replace(author.split(" ")[-1], "{author}")
            
            # Iterate over all other names of the author and remove them
            for name in author.split(" "):
                content = content.replace(name, "")

        except Exception as e:
            print(f"ERROR: Could not extract author from file!\n{e}")
            author = "unknown"
        
        # Get the filename (basename) and remove the file property
        title = self.extract_title(docx_path)
        
        # Initialize DocxDocument object
        print(f"{datetime.now().strftime("%H:%M:%S")}\t Create Docx document <{title[:20]}...> with content of length {len(content)}")
        docx_doc = DocumentHandler(docx_path, content, title)
        docx_doc.author = author
        
        docx_doc = self.process(docx_doc)
        
        return docx_doc
    
    def process_pdf(self, pdf_path: str) -> PdfDocument:
        """ Extracts the content from the pdf_path and process all questions

        Returns:
            PdfDocument: PDF-Document Handler with enriched informations.
        """        
        content = self.extract_text_from_pdf(pdf_path)
        content = self.clean_input(content, hard_clean=False, line_breaks=False)

        title = self.extract_title(pdf_path)
        
        # Initialize PdfDocument object
        print(f"{datetime.now().strftime("%H:%M:%S")}\t Create PDF document <{title[:20]}...> with content of length {len(content)}")
        pdf_doc = PdfDocument(pdf_path, content, title)
        
        pdf_doc = self.process(pdf_doc)
        
        return pdf_doc
    
    def create_wordcloud(self, doc: DocumentHandler | PdfDocument, wordcloud_names: list) -> DocumentHandler | PdfDocument:
        """ Creates the word clouds for one document. 
        
        Returns:
            DocumentHandler | PdfDocument: Document with enriched wordcloud data. """
            
        for wordcloud_name in wordcloud_names:
            path = os.path.join(self.output_folder, doc.file_name)

            if wordcloud_name == "markings":
                sentences = doc.marked_sentences
            else:
                content = doc.__dict__.get(wordcloud_name)
                sentences = content.split(".")

            new_wordcloud_data = self.wordcloud.process_wordcloud(input=sentences, path=path, wordcloud_name=wordcloud_name)
            doc.wordcloud_data.update(new_wordcloud_data)
        
        return doc
        
    
    def process_folder(self, files_folder: str, file_types: tuple) -> list[DocumentHandler | PdfDocument]:
        """ Iterates over all Files in the folder and processes files that match the file_types them.
        
        Args:
            files_folder (str): Source folder
            file_types (tuple): Filetypes that should be processed (e.g., .pdf, .docx)

        Returns:
            List[DocumentHandler | PdfDocument]: All documents that were processed.
        """
        all_documents = []
        for file_name in os.listdir(files_folder):
            if file_name.endswith(file_types):
                print(f"{datetime.now().strftime("%H:%M:%S")} Analyzing file from folder: {file_name}")
                doc_path = os.path.join(files_folder, file_name)
                
                if file_name.endswith("pdf"):
                    doc = self.process_pdf(doc_path)
                    
                elif file_name.endswith("docx"):
                    doc = self.process_docx(doc_path)
                    
                all_documents.append(doc)
                
                print("Temporarily storing documents")
                self.save_documents(self.all_documents)
                
        return all_documents
    
    def clean_input(self, input, hard_clean: bool=False, line_breaks: bool=False) -> str:
        """ Cleans the input. Applies to string, list and dictionaries.

        Args:
            input (str, list, dict): Input text
            hard_clean (bool): Whether line breaks and special characters should be removed. Defaults to False.
            line_breaks (bool): Whether line breaks should be considered. Defaults to False.

        Returns:
            str: Cleaned input files.
        """
        def remove_large_digits(input_string):
            """ Removes any digit that appears 3 or more times consecutively. """
            pattern = r'\d{3,}'  
            return re.sub(pattern, '', input_string)

        def remove_entries(words_to_remove: list, text: str):
            """ Removes entries from a text based on a list provided. """
            
            '''# Regex pattern to match whole words in the list
            pattern = r'\b(?:' + '|'.join(re.escape(word.lower()) for word in words_to_remove) + r')\b'
            
            # Replace matched words with an empty string
            cleaned_text = re.sub(pattern, '', text, flags=re.IGNORECASE)

            # Return cleaned text with excess spaces normalized
            return ' '.join(cleaned_text.split())'''
            def replacer(match):
                word = match.group(0)
                if word.lower() in words_to_remove:
                    return ''  # Remove the matched word
                return word  # Retain other words as-is

            # Use re.sub with the replacer function to remove words regardless of case sensitivity
            pattern = r'\b(?:' + '|'.join(re.escape(word) for word in words_to_remove) + r')\b'
            cleaned_text = re.sub(pattern, replacer, text, flags=re.IGNORECASE)

            # Normalize spaces and return cleaned text
            return ' '.join(cleaned_text.split())
            
        
        if isinstance(input, str):
            
            input = input.replace("’", "'") # Replace ’ with '
            
            input = re.sub(r"  +", " ", input) # Remove all double spaces
            
            input = input.replace("\t"," ") # Remove tabs

            # Remove line breaks if line_breaks is True and not hard_clean
            if line_breaks:
                input = input.replace("\n", " ")
            
            # Remove non-ascii characters if not hard_clean
            if hard_clean:
                words_to_remove = ["datum", "link", "dokument"]
                input = remove_entries(words_to_remove, input) # Remove words

                input = input.encode("ascii", "ignore").decode()  
                input = re.sub(r"[^a-zA-Z0-9.,*' -]", " ", input)
                input = remove_large_digits(input) # Remove large digits (usually years or bugs)
                 
            # Remove leading and trailing whitespaces
            input = input.strip()
            
            return input
        
        elif isinstance(input, list):
            # If input is a list, clean each element recursively
            return [self.clean_input(item, hard_clean, line_breaks) for item in input]
        
        elif isinstance(input, dict):
            # If input is a dictionary, clean each value recursively
            return {key: self.clean_input(value, hard_clean, line_breaks) for key, value in input.items()}
        
        else:
            # Return the input unchanged if it is not a string, list, or dictionary
            return input
    
    def save_documents(self, documents: list[DocumentHandler | PdfDocument]):
        """ Saves the processed documents as pickle file. """
        
        path = os.path.join(self.output_folder, self.file_name)
        
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
        
        with open(path, "wb") as f:
            pickle.dump([doc.to_dict() for doc in documents], f)
    
    def load_documents(self, input_path: str, load_latest=False):
        """ Loads the processed documents from a pickle file. """
        
        if os.path.exists(input_path):
            # Open the provided file
            with open(input_path, "rb") as f:
                data = pickle.load(f)
                
        # Load latest file for the model
        elif load_latest:
            # Get the list of all pickle files in the OUTPUT_FOLDER
            pkl_files = glob.glob(os.path.join(self.output_folder, "*.pkl"))
            if not pkl_files:
                raise FileNotFoundError("No pickle files found in the output folder.")

            # Find the latest pickle file based on the modification time
            latest_pkl_file = max(pkl_files, key=os.path.getmtime)

            # Load the latest pickle file
            with open(latest_pkl_file, "rb") as f:
                data = pickle.load(f)
        
        # Process the loaded data
        self.all_documents = [
            PdfDocument.from_dict(doc_dict) 
            if doc_dict.get("path", "").split(".")[-1] == "pdf" 
            else DocumentHandler.from_dict(doc_dict) 
            for doc_dict in data
        ]
    
    def export_docx_files(self):
        """ Exports all documents as docx files in output folder of the instance. """
        for doc in self.all_documents:
            file_name = f"cai_media_analysis_{doc.file_name}.docx"
            file_path = os.path.join(self.output_folder, file_name)
            doc.save_as_docx(file_path=file_path)
    
    def export_markdown_files(self):
        """ Exports all document as markdown files in output folder of the instance. """
        for doc in self.all_documents:
            # Create folders
            file_name = f"cai_media_analysis_{doc.file_name}.md"
            file_path = os.path.join(self.output_folder, file_name)
            
            # Write markdown file
            self.export_markdown(file_path, doc.get_markdown())
                
    def export_markdown(self, path: str, content: str):
        """ Export a specific content to the provided path. """
        with open(path, "w") as f:
            f.write(content)
                
    def __iter__(self, which:str ="all") -> list:
        """ Returns all or specific documents from the instance.

        Args:
            which (str): Title of the document that should be used to iterate over. Defaults to "all".

        Returns:
            list[PdfDocument]: List of document(s).
        """
        if which == "all":
            return iter(self.all_documents)
        else:
            for doc in self.all_documents:
                if doc.title == which:
                    return iter(doc)
                
                
if __name__ == "__main__":
    print("ToDo")